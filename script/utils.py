import os
import re
import docx
import json
from pysbd import Segmenter

def read_docx_files(directory):
    """读取文件"""
    docx_files = {}
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory, filename)
            doc = docx.Document(file_path)
            file_name_without_ext = os.path.splitext(filename)[0]
            docx_files[file_name_without_ext] = doc
    return docx_files

def process_docx_files_2_para(docx_files_dict):
    """文本分段处理"""
    docx_files_dict_processed = {}
    for filename, doc in docx_files_dict.items():
        docx_files_dict_processed[filename] = []
        current_paragraph = ''
        for para in doc.paragraphs:
            if len(para.text) == 0:
                continue
            if len(current_paragraph) + len(para.text) > 500:
                current_paragraph += para.text
                docx_files_dict_processed[filename].append(current_paragraph.strip())
                current_paragraph = para.text  # 保证上下文重叠
            else:
                current_paragraph += para.text  # 标记段落分割
        # Add the last accumulated paragraph
        if len(current_paragraph) > 100:
            docx_files_dict_processed[filename].append(current_paragraph.strip())
    return docx_files_dict_processed

def process_docx_files_2_sents(docx_files_dict):
    """文本分句处理"""
    docx_files_dict_processed_1 = {}
    docx_files_dict_processed_2 = {}
    for filename, doc in docx_files_dict.items():
        all_para = []
        for para in doc.paragraphs:
            all_para.append(para.text)
        docx_files_dict_processed_1[filename] = ''.join(all_para)

    left_list = [
        r'^（\d）',  # 句子开头：（1）
        r'^\d）',  # 句子开头：1）
        r'^\d\.\d',  # 句子开头：1.1
        r'^\d\.',  # 句子开头：1.
    ]
    right_list = [
        r'（\d）$',  # 句子末尾：（1）
        r'\d）$',  # 句子末尾：1）
        r'\d\.\d$'  # 句子末尾：1.1
        r'\d\.$',  # 句子末尾：1.
        r'\d\．$',  # 句子末尾：1．
        r'①.$',
        r'②.$',
        r'③.$',
    ]
    left_patterns = [re.compile(pattern) for pattern in left_list]
    right_patterns = [re.compile(pattern) for pattern in right_list]

    for filename, doc in docx_files_dict_processed_1.items():
        docx_files_dict_processed_2[filename] = []
        segmenter = Segmenter()
        para_2_sentences = segmenter.segment(doc)

        i = 0
        while i < len(para_2_sentences):
            sentence = para_2_sentences[i].replace(" ", "").replace("　", " ").replace("\u3000", " ").replace("\xa0", " ")
            # 跳过‘招标文件’最后三个句子
            name_list = ['招标', '定商', '承包', '采购', '货代', '服务', '油', '新能源']
            for name in name_list:
                if (filename.find(name) != -1) and (len(para_2_sentences) - i <= 6):
                    i += 1
                    continue
            # 跳过超短句
            if len(sentence.strip()) <= 3:
                i += 1
                continue
            # 检查并删除模式
            for pattern in left_patterns:
                match = pattern.search(sentence)
                if match:
                    sentence = sentence[match.end():].lstrip()
            for pattern in right_patterns:
                match = pattern.search(sentence)
                if match:
                    sentence = sentence[:match.start()].rstrip()
            # 连接下一句并跳过它
            if i + 1 < len(para_2_sentences) and para_2_sentences[i + 1].startswith('》'):
                sentence += para_2_sentences[i + 1]
                i += 1  # 跳过下一个句子
            docx_files_dict_processed_2[filename].append(sentence.strip())
            i += 1  # 处理下一个句子

    return docx_files_dict_processed_2

def process_docx_files_2_sents_WithNum(docx_files_dict):
    """文本分句处理"""
    docx_files_dict_processed_1 = {}
    docx_files_dict_processed_2 = {}
    for filename, doc in docx_files_dict.items():
        all_para = []
        for para in doc.paragraphs:
            all_para.append(para.text)
        docx_files_dict_processed_1[filename] = ''.join(all_para)

    for filename, doc in docx_files_dict_processed_1.items():
        docx_files_dict_processed_2[filename] = []
        segmenter = Segmenter()
        para_2_sentences = segmenter.segment(doc)

        i = 0
        while i < len(para_2_sentences):
            sentence = para_2_sentences[i].replace(" ", "").replace("　", " ").replace("\u3000", " ").replace("\xa0", " ")
            # 跳过‘招标文件’最后三个句子
            name_list = ['招标', '定商', '承包', '采购', '货代', '服务', '油', '新能源']
            for name in name_list:
                if (filename.find(name) != -1) and (len(para_2_sentences) - i <= 6):
                    i += 1
                    continue
            # 跳过超短句
            if len(sentence.strip()) <= 3:
                i += 1
                continue
            # 连接下一句并跳过它
            if i + 1 < len(para_2_sentences) and para_2_sentences[i + 1].startswith('》'):
                sentence += para_2_sentences[i + 1]
                i += 1  # 跳过下一个句子
            docx_files_dict_processed_2[filename].append(sentence.strip())
            i += 1  # 处理下一个句子

    return docx_files_dict_processed_2

def tender_document(docx_files_dict):
    """提取招标文件：资质要求部分"""
    docx_files_dict_processed = {}
    filename_list = ['招标', '定商', '承包', '采购', '货代', '变更', '油', '新能源']
    part_start_list = ['投标人资格要求', '申请人资格要求及技术要求', '投标人的资格要求']
    part_end_list = ['招标文件的获取', '集中资格招标文件的获取', '获取招标文件', '招标文件获取']

    for filename, doc in docx_files_dict.items():
        for name in filename_list:
            if filename.find(name) != -1:
                qualification_part = []
                start_flag = 0
                end_flag = 0
                for sentence in doc:
                    for pattern in part_start_list:
                        if sentence.find(pattern) != -1:
                            start_flag = 1
                    for pattern in part_end_list:
                        if sentence.find(pattern) != -1:
                            end_flag = 1
                    if start_flag == 1 and end_flag == 0:
                        qualification_part.append(sentence)
                    if start_flag == 1 and end_flag == 1:
                        break
                if len(qualification_part) > 0:
                    docx_files_dict_processed[filename] = qualification_part

    return docx_files_dict_processed

def tender_document_2_para(docx_files_dict):
    docx_files_dict_processed = {}
    word_list = ['要求', '申请人承诺', '质量管理', '财务状况', '禁止投标', '技术标准', '存在下列情形']
    for filename, doc in docx_files_dict.items():
        paras = []
        para = []
        for sentence in doc:
            if any(sentence.find(word) != -1 for word in word_list) and (len(para) != 0):
                paras.append(''.join(para))
                para = []
                para.append(sentence)
            else:
                para.append(sentence)
        if (len(para) != 0):
            paras.append(''.join(para))
        docx_files_dict_processed[filename] = paras
    return docx_files_dict_processed

def split_sentences(line):
    line_split = re.split(r'[。！；？]', line.strip())
    line_split = [line.strip() for line in line_split if
                  line.strip() not in ['。', '！', '？', '；', '，'] and len(line.strip()) > 1]
    return line_split

def process_docx_files_sents_2_longer(docx_files_dict, num):
    """文本分句合成更长分句"""
    docx_files_dict_processed = {}
    for filename, doc in docx_files_dict.items():
        docx_files_dict_processed[filename] = []
        i = 0
        while i < len(doc):
            combined_sentence = ''.join(doc[i:i + num])
            docx_files_dict_processed[filename].append(combined_sentence)
            i += num
    return docx_files_dict_processed

def process_docx_files_sents_2_para(docx_files_dict, para_size):
    """文本分句转分段"""
    docx_files_dict_processed = {}
    for filename, doc in docx_files_dict.items():
        docx_files_dict_processed[filename] = []
        current_paragraph = ''
        for sentence in doc:
            current_paragraph += sentence.strip()
            if len(current_paragraph) > para_size:
                docx_files_dict_processed[filename].append(current_paragraph.strip())
                current_paragraph = sentence.strip()
        if len(current_paragraph) > 100:
            docx_files_dict_processed[filename].append(current_paragraph.strip())
    return docx_files_dict_processed

def process_docx_files_year(docx_files_dict_processed):
    for filename, doc in docx_files_dict_processed.items():
        print(f"文件名: {filename}")
        print(f"句子数: {len(doc)}")
        for i, sentence in enumerate(doc):
            if i > 8:
                filename_list = ['招标', '定商', '承包', '采购', '货代', '变更', '油', '新能源']
                if any(filename.find(word) != -1 for word in filename_list):
                    year_str = '这篇招标文件发布于2024年，投标、开标截止日期在2024年范围内的一般是正确的，但也要注意可能的时间颠倒错误。'
                else:
                    year_str = '空'
                break
            match = re.search(r'\b20\d{2}\b', sentence)
            if match:
                year_str = sentence.strip()  # 保存符合条件的句子（去除首尾空白）
                break  # 找到符合条件的句子后退出内层循环
        doc.insert(0, year_str)
        print(year_str)
    return docx_files_dict_processed

def extract_context(i, doc, char_num):
    context_upper = ''  # 提取上文
    for j in range(i - 1, -1, -1):  # 从当前句子的前一句开始向前遍历
        context_upper = doc[j] + context_upper
        if len(context_upper) >= char_num:
            break
    context_lower = ''  # 提取下文
    for j in range(i + 1, len(doc)):  # 从当前句子的后一句开始向后遍历
        context_lower += doc[j]
        if len(context_lower) >= char_num:
            break
    # 返回整合的上下文
    return context_upper.strip(), context_lower.strip()

def clean_json_delimiters(text):
    text = text.strip()
    start_delimiter = '```json'
    end_delimiter = '```'
    # 检查开头是否有 '```json'
    if text.startswith(start_delimiter):
        text = text[len(start_delimiter):].lstrip()  # 删除开头的 '{```json}' 并去除空格
    # 检查结尾是否有 '```'
    if text.endswith(end_delimiter):
        text = text[:-len(end_delimiter)].rstrip()  # 删除结尾的 '{```}' 并去除空格
    # 检查结尾是否有 '<|im_end|>'
    qwen_end_delimiter = '<|im_end|>'
    if text.endswith(qwen_end_delimiter):
        text = text[:-len(qwen_end_delimiter)].rstrip()
    return text

if __name__ == '__main__':
    str = """
        ```json
        {
            "num": [
                "2"
            ],
            "error_sentence": [
                "一般招标项目不得采取随机抽取方式"
            ]
        }
        ```
    """
    parsed_json = json.loads(clean_json_delimiters(str))